from typing import Dict, List, Optional
from typing_extensions import TypedDict
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from dotenv import load_dotenv
import os
from IPython.display import Image, display, Markdown
from langgraph.graph import StateGraph, START, END
from langchain_community.tools import DuckDuckGoSearchResults
from langchain.agents import AgentExecutor, create_tool_calling_agent
from datetime import datetime
from langchain_google_genai import ChatGoogleGenerativeAI
import time
from pydantic import BaseModel
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import PyPDF2
from docx import Document as DocxDocument
load_dotenv()
from langchain.document_loaders import PDFMinerLoader

from arize.otel import register

from openinference.instrumentation.langchain import LangChainInstrumentor

tracer_provider = register(
    space_id = "U3BhY2U6MTYwNjc6Vmd2bA==",
    api_key = "5413c22fe01f8e49b37",
    project_name = "prepto", 
    set_global_tracer_provider=False
    )
    

LangChainInstrumentor().instrument(tracer_provider= tracer_provider)

# Initialize the language model
llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", api_key=os.environ['GEMINI_API_KEY'])

class RagAgent:
    def __init__(self, resume_content=None):
        self.model = llm
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=os.environ['GEMINI_API_KEY'])
        self.vector_store = None
        if resume_content:
            self.load_resume_into_vector_store(resume_content)
    def load_resume_into_vector_store(self,resume_content):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size  = 500,chunk_overlap = 50)
        chunks = text_splitter.split_text(resume_content)
        documents = [Document(page_content = chunk) for chunk in chunks]
        self.vector_store = FAISS.from_documents(documents , self.embeddings)
    def query_vector(self,query,k=3):
        if not self.vector_store:
            return "No resume content loaded"
        results = self.vector_store.similarity_search(query,k=k)
        return "\n".join([res.page_content for res in results])
    
    def grade_resume(self,state):
        if not self.vector_store:
            return "No resume available to Grade"
        resume_grade_prompt = ChatPromptTemplate.from_messages([
            ("system",'''You are a resume grading expert. Evaluate the resume content based on:
            - Clarity and structure
            - Relevance of skills and experience to the specified role
            - Inclusion of trending keywords and technologies
            - Strength of projects (if any)
            Provide a detailed evaluation in Markdown format with a score out of 10 for each criterion and an overall score.'''),
            ("human",f"Resume content :{self.query_vector('full resume content')}\n Target Role : {state['query']}")
            
        ])
        grading_meassages = resume_grade_prompt.invoke({"query":state['query']}).to_messages()
        grading_response = self.model.invoke(grading_meassages)
        path = save_with_format_choice(grading_response.content,"Resume_Grade")
        

    def generate_resume_based_question(self,num_questions = 3):
        if not self.vector_store:
            return ["No resume available, kindly upload your resume"]
        project_content = self.query_vector("projects")
        tech_stack = self.query_vector("skills OR technologies OR tech stack")
        question_prompt = ChatPromptTemplate.from_messages(
            [("system", '''You are an expert interviewer. Based on the provided resume excerpts, generate {num_questions} specific, concise interview questions focusing on:
            - Projects mentioned in the resume
            - Technologies and skills listed
            Return the questions as a numbered list.'''),("human",f"Projects:{project_content}\nTech Stack:{tech_stack}")]

        )

        question_messages = question_prompt.invoke({"num_questions":num_questions})
        question_response = self.model.invoke(question_messages)
        questions = [q.strip() for q in question_response.content.split("\n") if q.strip() and q[0].isdigit()]
        return questions        
        



# Define the state structure with chat history and additional fields
class State(TypedDict):
    query: str
    response: Optional[str]
    category: Optional[str]
    chat_history: List  # List of HumanMessage and AIMessage objects
    resume_details: Optional[Dict[str, str]]  # For resume creation
    interview_responses: Optional[List[str]]  # For mock interview responses
    start_time: Optional[float]  # For timing the mock interview
    interview_details: Optional[Dict[str, str]]  # For interview question clarification
    job_search_details: Optional[Dict[str, str]]  # For job search clarification
    learning_details: Optional[Dict[str, str]]  # For learning resource clarification
    rag_agent: Optional[RagAgent]
    resume_path: Optional[str]

# Save content to a file
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def save_file(data, filename, format_choice="md"):
    folder_name = "Agent_output"
    os.makedirs(folder_name, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = f"{filename}_{timestamp}"
    file_path = os.path.join(folder_name, f"{file_name}.{format_choice}")

    if format_choice == "md":
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(data)
    elif format_choice == "docx":
        doc = DocxDocument()
        for line in data.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(file_path)
    elif format_choice == "pdf":
        pdf = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for line in data.split('\n'):
            if line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.strip():
                story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 12))
        pdf.build(story)
    else:
        raise ValueError("Unsupported format. Choose 'md', 'docx', or 'pdf'.")

    print(f"File {file_path} created successfully")
    return file_path

def save_with_format_choice(data, filename):
    print("\nChoose output format:\n1. Markdown (.md)\n2. Word Document (.docx)\n3. PDF (.pdf)")
    while True:
        choice = input("Enter 1, 2, or 3: ").strip()
        format_map = {"1": "md", "2": "docx", "3": "pdf"}
        format_choice = format_map.get(choice)
        if format_choice:
            break
        print("Invalid choice. Please enter 1, 2, or 3.")
    return save_file(data, filename, format_choice)



# Display file content as Markdown
def show_file(file_path):
    if not isinstance(file_path, str) or not os.path.exists(file_path):
        print(f"Cannot display file: {file_path}")
        return
    
    file_extension = os.path.splitext(file_path.lower())[1]
    if file_extension in [".docx", ".pdf"]:
        print(f"File saved as {file_path}. Open it with a compatible viewer (e.g., Word for .docx, Acrobat for .pdf).")
        return
    
    # Only try to read and display text-based files
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        display(Markdown(content))
    except UnicodeDecodeError:
        print(f"Cannot display {file_path}: File is not in a readable text format.")
    except Exception as e:
        print(f"Error displaying {file_path}: {str(e)}")

# Learning Resource Agent
class LearningResourceAgent:
    def __init__(self, prompt):
        self.model = llm
        self.prompt = prompt
        self.tools = [DuckDuckGoSearchResults()]

    def TutorialAgent(self, state):
        agent = create_tool_calling_agent(self.model, self.tools, self.prompt)
        agent_executor = AgentExecutor(agent=agent, tools=self.tools, verbose=True)
        response = agent_executor.invoke({
            "input": state['query'],
            "chat_history": state.get('chat_history', [])
        })
        path = save_with_format_choice(str(response.get('output', '')).replace("```markdown", "").strip(), 'Tutorial')
        return path
    
    def QueryBot(self, state):
        if 'learning_details' not in state or not state['learning_details']:
            state['learning_details'] = {}

        # Generate dynamic clarification questions based on the initial query
        clarification_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert in guiding users to learning resources. Based on the user's initial query, generate 3 concise, relevant clarification questions to tailor the recommendations (e.g., about current knowledge level, preferred format, or goals). Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        clarification_messages = clarification_prompt.invoke({"query": state['query']}).to_messages()
        clarification_response = self.model.invoke(clarification_messages)
        clarification_questions = [q.strip() for q in clarification_response.content.split('\n') if q.strip() and q[0].isdigit()]

        # Collect answers for the clarification questions
        for i, question in enumerate(clarification_questions[:3], 1):
            if f"answer_{i}" not in state['learning_details']:
                print(f"\nAI: {question}")
                answer = input("Your answer: ")
                state['learning_details'][f"answer_{i}"] = answer if answer.strip() else "None specified"

        # Generate tailored recommendations with general ones included
        tailored_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert in providing learning resources. Based on the user's initial query and their clarification answers, generate tailored recommendations for learning sources. Include the following general recommendations after your tailored suggestions, categorized as shown:

Return the response in Markdown format with sections: ## Tailored Recommendations, ## General Recommendations (including all the above).'''),
            ("human", f"Initial query: {state['query']}\nClarification 1: {state['learning_details']['answer_1']}\nClarification 2: {state['learning_details']['answer_2']}\nClarification 3: {state['learning_details']['answer_3']}")
        ])
        tailored_messages = tailored_prompt.invoke({"query": state['query']}).to_messages()
        tailored_response = self.model.invoke(tailored_messages)
        
        path = save_with_format_choice(tailored_response.content, 'Q&A_Doubt_Session')
        return path

def extract_resume(file_path):
    if not os.path.exists(file_path):
        return {"status": "error", "message": "File not found"}
    file_extension = os.path.splitext(file_path.lower())[1]
    try:
        if file_extension == ".pdf":
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return {"status": "success", "content": text}
        elif file_extension in ['.doc', ".docx"]:
            doc = DocxDocument(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return {"status": "success", "content": text}
        else:
            return {"status": "error", "message": "Unsupported file format. Please upload a PDF or DOC/DOCX file."}
    except Exception as e:
        return {"status": "error", "message": f"Error reading file: {str(e)}"}



# Interview Agent
class InterviewAgent:
    def __init__(self, prompt,rag_agent = None):
        self.model = llm
        self.prompt = prompt
        self.tools = [DuckDuckGoSearchResults()]
        self.rag_agent = rag_agent
    
    def Interview_questions(self, state):
        if 'interview_details' not in state or not state['interview_details']:
            state['interview_details'] = {}

        clarification_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled interviewer tasked with generating clarification questions to tailor a list of interview questions. Based on the user's initial query, generate 2 concise, relevant clarification questions to refine the list (e.g., about experience level, specific areas of interest, or role specifics). Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        clarification_messages = clarification_prompt.invoke({"query": state['query']}).to_messages()
        clarification_response = self.model.invoke(clarification_messages)
        clarification_questions = [q.strip() for q in clarification_response.content.split('\n') if q.strip() and q[0].isdigit()]

        for i, question in enumerate(clarification_questions[:2], 1):
            if f"answer_{i}" not in state['interview_details']:
                print(f"\nAI: {question}")
                answer = input("Your answer: ")
                state['interview_details'][f"answer_{i}"] = answer if answer.strip() else "None specified"

        questions_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled researcher tasked with generating interview questions for a role based on the user's initial query and additional details provided. Generate exactly 10 relevant interview questions in the following order:
- 5 Basic Technical Questions ( foundational concepts suitable for beginners)
- 5 Intermediate Technical Questions ( moderately complex topics requiring some experience include intermediate coding problems)
- 5 Advanced Technical Questions ( complex, in-depth topics for experienced candidates , include advance coding problems)
- 5 Behavioral Question ( focused on past experience or teamwork and projects )
Tailor the questions to the provided details. Return the questions in Markdown format with clear section headers in this order: ## Basic Technical Questions, ## Intermediate Technical Questions, ## Advanced Technical Questions, ## Behavioral Question.'''),
            ("human", f"Initial query: {state['query']}\nClarification 1: {state['interview_details']['answer_1']}\nClarification 2: {state['interview_details']['answer_2']}")
        ])
        
        question_messages = questions_prompt.invoke({"query": state['query']}).to_messages()
        question_response = self.model.invoke(question_messages)
        
        path = save_with_format_choice(question_response.content, 'Interview_Questions')
        return path
    
    def Mock_Interview(self, state):
        if 'start_time' not in state or not state['start_time']:
            state['start_time'] = time.time()
        if 'interview_responses' not in state or not state['interview_responses']:
            state['interview_responses'] = []
        
        max_duration = 60 * 60  # 60 minutes in seconds
        print("Session started.. \nYou will be given a total of 20 questions:\n- 5 based on your resume\n- 5 Basic technical questions\n- 5 Intermediate technical questions\n- 5 Advanced technical questions\nYou have 60 minutes..")

        # Check if resume is uploaded, prompt if not
        resume_questions = []
        if state.get('rag_agent'):
            print("Using previously uploaded resume for questions...")
            resume_questions = state['rag_agent'].generate_resume_based_question(num_questions=5)
        else:
            print("\nNo resume uploaded yet. Please provide the file path to your resume (PDF or DOC/DOCX) to generate resume-based questions.")
            file_path = input("File path: ")
            resume_result = extract_resume(file_path)
            
            if resume_result["status"] == "error":
                print(f"Error: {resume_result['message']}. Using placeholder questions instead.")
                resume_questions = ["No resume available - placeholder question"] * 5
            else:
                resume_content = resume_result["content"]
                state['rag_agent'] = RagAgent(resume_content=resume_content)
                state['resume_path'] = file_path
                print("Resume uploaded successfully. Generating resume-based questions...")
                resume_questions = state['rag_agent'].generate_resume_based_question(num_questions=5)

        # Generate remaining 15 questions
        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert interviewer tasked with conducting a mock interview. Based on the user's initial query, generate 15 concise, relevant interview questions tailored to the specified role or topic:
    - 5 Basic Technical Questions (foundational concepts suitable for beginners)
    - 5 Intermediate Technical Questions (moderately complex topics requiring some experience, include intermediate coding problems)
    - 5 Advanced Technical Questions (complex, in-depth topics for experienced candidates, include advanced coding problems)'''),
            ("human", f"Initial query: {state['query']}")
        ])
        question_messages = question_prompt.invoke({"query": state['query']}).to_messages()
        question_response = self.model.invoke(question_messages)
        other_questions = [q.strip() for q in question_response.content.split('\n') if q.strip() and q[0].isdigit()]
        
        # Combine questions
        all_questions = resume_questions + other_questions[:15]
        
        record = [f"Mock Interview for: {state['query']}\n\n"]
        
        for i, q in enumerate(all_questions[:20]):
            if time.time() - state['start_time'] > max_duration:
                record.append("Time limit reached. Ending the interview.\n")
                break
            print(f"\nInterviewer: {q}")
            response = input("Your answer: ")
            state['interview_responses'].append(response)
            record.append(f"Question {i+1}: {q}\nYour Response: {response}\n\n")
        
        responses_text = chr(10).join([f"Q{i+1}: {q}\nA: {r}" for i, (q, r) in enumerate(zip(all_questions[:len(state['interview_responses'])], state['interview_responses']))])
        eval_prompt = ChatPromptTemplate.from_messages([
            ("system", self.prompt[0].content),
            ("human", "Evaluate these mock interview responses:" + chr(10) + responses_text)
        ])
        eval_messages = eval_prompt.invoke({}).to_messages()
        eval_response = self.model.invoke(eval_messages)
        record.append(f"AI Evaluation:\n{eval_response.content}\n\n")
        
        print("\nInterviewer: Thank you for completing the mock interview. Could you provide feedback on how you felt about this session (e.g., question difficulty, pacing, usefulness)?")
        candidate_feedback = input("Your feedback: ")
        record.append(f"Candidate Feedback:\n{candidate_feedback}\n")
        
        path = save_with_format_choice(''.join(record), 'Mock_Interview')
        return path


# Resume Maker
class ResumeMaker:
    def __init__(self, prompt):
        self.model = llm
        self.prompt = prompt
        self.tools = [DuckDuckGoSearchResults()]
        self.agent_executor = AgentExecutor(
            agent=create_tool_calling_agent(self.model, self.tools, self.prompt),
            tools=self.tools,
            verbose=True
        )
        self.rag_agent = None
    
    def Create_Resume(self, state):
        if 'resume_details' not in state or not state['resume_details']:
            state['resume_details'] = {}

        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a resume expert tasked with gathering details from a user to create a resume. Based on the user's initial query, generate 5 concise, relevant questions to collect essential information (e.g., role, experience, skills, projects, education) for a professional resume. Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        question_messages = question_prompt.invoke({"query": state['query']}).to_messages()
        question_response = self.model.invoke(question_messages)
        questions = [q.strip() for q in question_response.content.split('\n') if q.strip() and q[0].isdigit()]
        
        for i, question in enumerate(questions[:5], 1):
            print(f"\nAI: {question}")
            answer = input("Your answer: ")
            state['resume_details'][f"answer_{i}"] = answer
        
        resume_content = f"# Resume\n\nBased on: {state['query']}\n\n"
        for i, answer in state['resume_details'].items():
            resume_content += f"**Detail {i.split('_')[1]}**: {answer}\n\n"
        
        refine_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled resume expert. Refine this raw resume content into a polished, professional template in Markdown format, incorporating trending keywords and technologies relevant to the specified role.'''),
            ("human", f"Raw resume content:\n{resume_content}")
        ])
        refined_messages = refine_prompt.invoke({"input": state['query']}).to_messages()
        response = self.model.invoke(refined_messages)
        path = save_with_format_choice(response.content, 'Resume')
        return path

# Job Search
class JobSearch:
    def __init__(self, prompt):
        self.model = llm
        self.prompt = prompt
        self.tools = DuckDuckGoSearchResults()
    
    def find_jobs(self, state):
        if 'job_search_details' not in state or not state['job_search_details']:
            state['job_search_details'] = {}

        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a job search assistant...'''),
            ("human", f"Initial query: {state['query']}")
        ])
        question_messages = question_prompt.invoke({"query": state['query']}).to_messages()
        question_response = self.model.invoke(question_messages)
        clarification_question = question_response.content.strip()

        print(f"\nAI: {clarification_question}")
        answer = input("Your answer: ")
        state['job_search_details']['clarification'] = answer

        refined_query = f"{state['query']} {state['job_search_details']['clarification']}"
        #print(f"DEBUG: Refined query is: {refined_query}")
        try:
            results = self.tools.run(refined_query)
            if not results:
                results = "No results found for this query. Try refining your search terms."
        except Exception as e:
            results = f"Search failed: {str(e)}. Please try again later."
        
        chain = self.prompt | self.model
        jobs = chain.invoke({"result": results}).content
        path = save_with_format_choice(str(jobs).replace("```markdown", "").strip(), 'Job_Search')
        return path

# Categorization Functions
def categories(state: State) -> State:
    if state["query"].lower() == "exit":
        return {"category": "exit"}
    prompt = ChatPromptTemplate.from_template(
        "Categorize the following customer query into one of these categories:\n"
        "1: Learn a specific Topic or Technology \n"
        "2: Resume Making\n"
        "3: Interview Preparation\n"
        "4: Job Search\n"
        "Give the number only as an output.\n\n"
        "Examples:\n"
        "1. Query: 'What are the basics of generative AI, and how can I start learning it?' -> 1\n"
        "2. Query: 'Can you help me improve my resume for a tech position?' -> 2\n"
        "3. Query: 'What are some common questions asked in AI interviews?' -> 3\n"
        "4. Query: 'Are there any job openings for AI engineers?' -> 4\n\n"
        "Now, categorize the following customer query:\n"
        "Query: {query}"
    )
    chain = prompt | llm
    print('categorizing the user query..')
    category = chain.invoke({"query": state['query']}).content
    state['chat_history'].append(HumanMessage(content=state['query']))
    return {"category": category}

def handle_learning_resources(state: State) -> State:
    prompt = ChatPromptTemplate.from_template(
        "Categorize the following user query into one of these categories:\n\n"
        "Categories:\n"
        "- Tutorial: For queries related to creating or requesting tutorials, guides, or documentation as learning resources for a specific topic.\n"
        "- Question: For general queries asking about learning resources or topics.\n"
        "- Default to Question if the query doesn't fit either of these categories.\n\n"
        "Examples:\n"
        "1. User query: 'How to create a tutorial on learning Python programming?' -> Category: Tutorial\n"
        "2. User query: 'Can you provide a step-by-step guide on mastering data structures?' -> Category: Tutorial\n"
        "3. User query: 'Provide me the documentation for learning cloud computing?' -> Category: Tutorial\n"
        "4. User query: 'What are the best resources for learning web development?' -> Category: Question\n"
        "5. User query: 'Are there any free courses available for machine learning?' -> Category: Question\n\n"
        "Now, categorize the following user query:\n"
        "The user query is: {query}\n"
    )
    chain = prompt | llm
    category = chain.invoke({"query": state['query']}).content
    return {'category': category}

def handle_interview_prep(state: State) -> State:
    prompt = ChatPromptTemplate.from_template(
        "Categorize the following user query into one of these categories:\n\n"
        "Categories:\n"
        "- Mock: For requests related to mock interviews on a user-specified topic.\n"
        "- Question: For general queries asking about interview topics or preparation.\n"
        "- Default to Question if the query doesn't fit either of these categories.\n\n"
        "Examples:\n"
        "1. User query: 'Can you conduct a mock interview with me for a software engineering role?' -> Category: Mock\n"
        "2. User query: 'What topics should I prepare for a data science interview?' -> Category: Question\n"
        "3. User query: 'I need to practice an interview focused on cloud computing.' -> Category: Mock\n"
        "4. User query: 'Can you list important coding topics for technical interviews?' -> Category: Question\n\n"
        "Now, categorize the following user query:\n"
        "The user query is: {query}\n"
    )
    chain = prompt | llm
    category = chain.invoke({"query": state["query"]}).content
    return {"category": category}

def job_search(state: State) -> State:
    prompt = ChatPromptTemplate.from_template(
        '''Your task is to refactor and make .md file for this content which includes
        the jobs available in the market. Refactor such that user can refer easily. Content: {result}'''
    )
    job_search_instance = JobSearch(prompt)
    path = job_search_instance.find_jobs(state)
    show_file(path)
    state['chat_history'].append(AIMessage(content=path))
    return {"response": path}

def handle_resume(state: State) -> State:
    prompt = ChatPromptTemplate.from_messages([
        ("system", '''You are a skilled resume expert with extensive experience in crafting resumes tailored for tech roles across various fields. 
Your task is to either create a resume template or grade an uploaded resume for a user-specified tech role or specialization. 
If creating, incorporate trending keywords and technologies relevant to the current job market for that topic and ask for details step-by-step (4-5 steps). 
If grading, use the uploaded resume content. Ensure outputs are in .md format.'''),
        MessagesPlaceholder('chat_history'),
        ("human", "{input}"),
        MessagesPlaceholder("agent_scratchpad")
    ])
    
    resume_maker = ResumeMaker(prompt)
    
    if "grade" in state['query'].lower():
        print("\nAI: Please provide the file path to your resume (PDF or DOC/DOCX):")
        file_path = input("File path: ")
        resume_result = extract_resume(file_path)
        
        if resume_result["status"] == "error":
            error_message = resume_result["message"]
            state['chat_history'].append(AIMessage(content=error_message))
            return {"response": error_message}
        
        # Resume extracted successfully
        resume_content = resume_result["content"]
        rag_agent = RagAgent(resume_content=resume_content)
        state['rag_agent'] = rag_agent
        state['resume_path'] = file_path
        
        # Grade the resume
        grade_result = rag_agent.grade_resume(state)
        if grade_result == "No resume available to Grade":
            error_message = "Failed to load resume content for grading."
            state['chat_history'].append(AIMessage(content=error_message))
            return {"response": error_message}
        
        # Since grade_result is currently just a string from query_vector, we need to generate the actual grading
        resume_grade_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a resume grading expert. Evaluate the resume content based on:
            - Clarity and structure
            - Relevance of skills and experience to the specified role
            - Inclusion of trending keywords and technologies
            - Strength of projects (if any)
            Provide a detailed evaluation in Markdown format with a score out of 10 for each criterion and an overall score.'''),
            ("human", f"Resume content: {resume_content}\nTarget Role: {state['query'].replace('upload resume for', '').strip()}")
        ])
        grading_messages = resume_grade_prompt.invoke({"query": state['query']}).to_messages()
        grading_response = rag_agent.model.invoke(grading_messages)
        
        # Save and show the grade with format choice
        grade_path = save_with_format_choice(grading_response.content, "Resume_Grade")
        show_file(grade_path)
        
        state['chat_history'].append(AIMessage(content=f"Uploaded resume graded: {grade_path}"))
        return {"response": f"Grade: {grade_path}"}
    
    else:
        # Resume creation logic
        resume_maker.rag_agent = RagAgent()
        resume_path = resume_maker.Create_Resume(state)
        show_file(resume_path)
        grade_path = resume_maker.rag_agent.grade_resume(state)
        if grade_path != "No resume available to Grade":  # Only show if it’s a valid path
            show_file(grade_path)
        
        state['chat_history'].append(AIMessage(content=f"Resume: {resume_path}\nGrade: {grade_path}"))
        state['rag_agent'] = resume_maker.rag_agent
        return {"response": f"Resume: {resume_path}\nGrade: {grade_path}"}
    
def query_bot(state: State) -> State:
    system_message = """You are an expert in a user-specified career field with extensive experience in training and guiding others in that domain. 
You have a strong track record of solving complex problems and addressing various challenges related to the specified career topic. 
Your role is to assist users by providing insightful solutions and expert advice on their queries about their chosen career field. 
Engage in a back-and-forth chat session to address user queries. 
Note: The user must specify the career topic (e.g., 'Software Engineering,' 'Data Analysis,' 'Cybersecurity,' or any other career field) for which they seek guidance."""
    prompt = [SystemMessage(content=system_message)]
    learning_agent = LearningResourceAgent(prompt)
    path = learning_agent.QueryBot(state)
    show_file(path)
    state['chat_history'].append(AIMessage(content=path))
    return {"response": path}

def tutorial_agent(state: State) -> State:
    system_message = """You are a knowledgeable assistant specializing as a Senior Developer with extensive experience in both development and tutoring for a user-specified topic. 
                            Additionally, you are an experienced blogger who creates tutorials focused on the topic provided by the user. 
                            Your task is to develop high-quality tutorial blogs in .md format with coding examples based on the user's requirements. 
                            Ensure the tutorial includes clear explanations, well-structured Python code, comments, and fully functional code examples. 
                            Provide resource reference links at the end of each tutorial for further learning. 
                            Note: The user must specify the topic (e.g., 'Web Development,' 'Data Science,' 'Machine Learning,' or any other field) for the tutorial blog."""
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_message),
        MessagesPlaceholder('chat_history'),
        ("human", "{input}"),
        MessagesPlaceholder("agent_scratchpad")
    ])
    learning_agent = LearningResourceAgent(prompt)
    path = learning_agent.TutorialAgent(state)
    show_file(path)
    state['chat_history'].append(AIMessage(content=path))
    return {"response": path}

def interview_topics_question(state: State) -> State:
    system_message = """You are a skilled researcher in finding interview questions for topics and jobs specified by the user. 
                            Your task is to provide a list of interview questions for the user-specified topic and job based on their requirements. 
                            Provide top questions with references and links if possible. May ask for clarification if needed. 
                            Generate a .md document containing the questions. 
                            Note: The user must specify the topic (e.g., 'Software Engineering,' 'Data Science,' 'Cybersecurity,' or any other field) for the interview questions."""
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_message),
        MessagesPlaceholder('chat_history'),
        ("human", "{input}"),
        MessagesPlaceholder("agent_scratchpad")
    ])
    interview_agent = InterviewAgent(prompt)
    path = interview_agent.Interview_questions(state)
    show_file(path)
    state['chat_history'].append(AIMessage(content=path))
    return {"response": path}

def mock_interview(state: State) -> State:
    system_message = """You are an expert interviewer for a user-specified role or topic. You have conducted numerous interviews for positions related to the topic provided by the user. 
                         Your task is to conduct a mock interview for the user-specified position or topic, engaging in a back-and-forth interview session. 
                         The conversation should not exceed more than 60 minutes. 
                         At the end of the interview, provide an evaluation for the candidate. 
                         Note: The user must specify the position or topic (e.g., 'Software Engineering,' 'Data Analysis,' 'Cloud Architecture,' or any other field) for the mock interview.If the user did not specify the role , kindly ask for the role before starting the interview."""
    prompt = [SystemMessage(content=system_message)]
    interview_agent = InterviewAgent(prompt)
    path = interview_agent.Mock_Interview(state)
    show_file(path)
    state['chat_history'].append(AIMessage(content=path))
    return {"response": path}

# Routing Functions
def route_query(state: State) -> str:
    if state["category"] == "exit":
        return END
    if '1' in state["category"]:
        print('Category: handle_learning_resource')
        return "handle_learning_resource"
    elif '2' in state["category"]:
        print('Category: handle_resume_making')
        return "handle_resume_making"
    elif '3' in state["category"]:
        print('Category: handle_interview_preparation')
        return "handle_interview_prep"
    elif '4' in state["category"]:
        print('Category: job_search')
        return "job_search"
    else:
        print("Please ask your question based on my description.")
        return END

def route_interview(state: State) -> str:
    if 'Question'.lower() in state["category"].lower():
        print('Category: interview_topics_question')
        return "interview_topics_question"
    elif 'Mock'.lower() in state["category"].lower():
        print('Category: mock_interview')
        return "mock_interview"
    else:
        print('Category: mock_interview')
        return "mock_interview"

def route_learning(state: State) -> str:
    if 'Question'.lower() in state["category"].lower():
        print('Category: query_bot')
        return "query_bot"
    elif 'Tutorial'.lower() in state["category"].lower():
        print('Category: tutorial_agent')
        return "tutorial_agent"
    else:
        print("Please ask your question based on my description.")
        return "query_bot"

# Build the Graph
graph = StateGraph(State)
graph.add_node("categories", categories)
graph.add_node("handle_learning_resource", handle_learning_resources)
graph.add_node("handle_resume_making", handle_resume)
graph.add_node("handle_interview_prep", handle_interview_prep)
graph.add_node("job_search", job_search)
graph.add_node("mock_interview", mock_interview)
graph.add_node("interview_topics_question", interview_topics_question)
graph.add_node("tutorial_agent", tutorial_agent)
graph.add_node("query_bot", query_bot)

graph.add_edge(START, "categories")
graph.add_conditional_edges("categories", route_query)
graph.add_conditional_edges("handle_interview_prep", route_interview)
graph.add_conditional_edges("handle_learning_resource", route_learning)

# Route handler nodes to END
for node in ["handle_resume_making", "job_search", "mock_interview", 
             "interview_topics_question", "tutorial_agent", "query_bot"]:
    graph.add_edge(node, END)


app = graph.compile()

# Conversational Loop
def run_conversation():
    print("Welcome! I'm here to assist you. You can:\n- Create a resume\n- Upload a resume for grading (say 'upload resume' in your query)\n- Conduct a mock interview (Please mention the role in your query)\n- Job Search\n- Interview topic questions\nType 'exit' to end the session.")
    # Initialize state once and reuse it
    state = {"query": "", "response": None, "category": None, "chat_history": [], 
             "resume_details": None, "interview_responses": None, "start_time": None, 
             "interview_details": None, "job_search_details": None, "learning_details": None, 
             "rag_agent": None, "resume_path": None}
    
    while True:
        state["query"] = input("\nYour query: ")
        if state["query"].lower() == "exit":
            print("Goodbye!")
            break
        results = app.invoke(state)
        state.update(results)  # Merge results into existing state instead of replacing it
        print(f"Response saved at: {state['response']}")
        print("\nConversation so far:")
        for entry in state["chat_history"]:
            role = "User" if isinstance(entry, HumanMessage) else "AI"
            print(f"{role}: {entry.content}")
        # Reset only temporary fields, keep rag_agent and resume_path
        state["response"] = None
        state["resume_details"] = None
        state["interview_responses"] = None
        state["start_time"] = None
        state["interview_details"] = None
        state["job_search_details"] = None
        state["learning_details"] = None


if __name__ == "__main__":
    run_conversation()