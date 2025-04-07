from fastapi import FastAPI, Request, Form, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import Dict, List, Optional
from typing_extensions import TypedDict
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from dotenv import load_dotenv
import os
import re
from langgraph.graph import StateGraph, START, END
from langchain_community.tools import DuckDuckGoSearchResults
from langchain.agents import AgentExecutor, create_tool_calling_agent
from datetime import datetime
from langchain_google_genai import ChatGoogleGenerativeAI
import time
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import PyPDF2
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import asyncio
from uuid import uuid4

load_dotenv()

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", api_key=os.environ['GEMINI_API_KEY'])

class State(TypedDict):
    query: str
    response: Optional[str]
    category: Optional[str]
    chat_history: List
    resume_details: Optional[Dict[str, str]]
    interview_responses: Optional[List[str]]
    start_time: Optional[float]
    interview_details: Optional[Dict[str, str]]
    job_search_details: Optional[Dict[str, str]]
    learning_details: Optional[Dict[str, str]]
    rag_agent: Optional['RagAgent']
    resume_path: Optional[str]
    current_step: Optional[int]
    questions: Optional[List[str]]
    content: Optional[str]
    instructions: Optional[str]

class RagAgent:
    def __init__(self, resume_content=None):
        self.model = llm
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=os.environ['GEMINI_API_KEY'])
        self.vector_store = None
        if resume_content:
            self.load_resume_into_vector_store(resume_content)

    def load_resume_into_vector_store(self, resume_content):
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(resume_content)
        documents = [Document(page_content=chunk) for chunk in chunks]
        self.vector_store = FAISS.from_documents(documents, self.embeddings)

    def query_vector(self, query, k=3):
        if not self.vector_store:
            return "No resume content loaded"
        results = self.vector_store.similarity_search(query, k=k)
        return "\n".join([res.page_content for res in results])

    def grade_resume(self, state):
        if not self.vector_store:
            return "No resume available to grade"
        resume_grade_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a resume grading expert. Evaluate the resume content based on:
            - Clarity and structure
            - Relevance of skills and experience to the specified role
            - Inclusion of trending keywords and technologies
            - Strength of projects (if any)
            Provide a detailed evaluation in Markdown format with a score out of 10 for each criterion and an overall score.'''),
            ("human", f"Resume content: {self.query_vector('full resume content')}\nTarget Role: {state['query'].replace('upload resume for', '').strip()}")
        ])
        grading_response = self.model.invoke(resume_grade_prompt.invoke({"query": state['query']}).to_messages())
        return grading_response.content

    def generate_resume_based_question(self, num_questions=3):
        if not self.vector_store:
            return ["No resume available, kindly upload your resume"] * num_questions
        project_content = self.query_vector("projects")
        tech_stack = self.query_vector("skills OR technologies OR tech stack")
        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert interviewer. Based on the provided resume excerpts, generate {num_questions} specific, concise interview questions focusing on:
            - Projects mentioned in the resume
            - Technologies and skills listed
            Return the questions as a numbered list.'''),
            ("human", f"Projects: {project_content}\nTech Stack: {tech_stack}")
        ])
        question_response = self.model.invoke(question_prompt.invoke({"num_questions": num_questions}).to_messages())
        return [q.strip() for q in question_response.content.split("\n") if q.strip() and q[0].isdigit()][:num_questions]

class InterviewAgent:
    def __init__(self, rag_agent=None):
        self.model = llm
        self.tools = [DuckDuckGoSearchResults()]
        self.rag_agent = rag_agent

    def generate_clarification_questions(self, state):
        clarification_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled interviewer tasked with generating clarification questions to tailor a list of interview questions or a mock interview. Based on the user's initial query, generate 2 concise, relevant clarification questions to refine the process (e.g., about experience level, specific areas of interest, or role specifics). Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        clarification_response = self.model.invoke(clarification_prompt.invoke({"query": state['query']}).to_messages())
        return [q.strip() for q in clarification_response.content.split('\n') if q.strip() and q[0].isdigit()][:2]

    def prepare_interview_questions(self, state):
        questions_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled researcher tasked with generating interview questions for a role based on the user's initial query and additional details provided. Generate exactly 15 relevant interview questions in the following order:
- 5 Basic Technical Questions (foundational concepts suitable for beginners)
- 5 Intermediate Technical Questions (moderately complex topics requiring some experience, include intermediate coding problems)
- 5 Advanced Technical Questions (complex, in-depth topics for experienced candidates, include advanced coding problems)
Tailor the questions to the provided details. Return the questions in Markdown format with clear section headers: ## Basic Technical Questions, ## Intermediate Technical Questions, ## Advanced Technical Questions.'''),
            ("human", f"Initial query: {state['query']}\nClarification 1: {state['interview_details'].get('answer_1', 'N/A')}\nClarification 2: {state['interview_details'].get('answer_2', 'N/A')}")
        ])
        question_response = self.model.invoke(questions_prompt.invoke({"query": state['query']}).to_messages())
        return question_response.content

    def mock_interview_questions(self, state):
        resume_questions = self.rag_agent.generate_resume_based_question(5) if self.rag_agent else ["No resume uploaded - generic question"] * 5
        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert interviewer tasked with generating questions for a mock interview. Based on the user's initial query, generate 15 concise, relevant interview questions tailored to the specified role or topic:
- 5 Basic Technical Questions (foundational concepts suitable for beginners)
- 5 Intermediate Technical Questions (moderately complex topics requiring some experience, include intermediate coding problems)
- 5 Advanced Technical Questions (complex, in-depth topics for experienced candidates, include advanced coding problems)'''),
            ("human", f"Initial query: {state['query']}\nClarification 1: {state['interview_details'].get('answer_1', 'N/A')}\nClarification 2: {state['interview_details'].get('answer_2', 'N/A')}")
        ])
        question_response = self.model.invoke(question_prompt.invoke({"query": state['query']}).to_messages())
        other_questions = [q.strip() for q in question_response.content.split('\n') if q.strip() and q[0].isdigit()]
        return resume_questions + other_questions[:15]

    def evaluate_mock_interview(self, state):
        record = [f"Mock Interview for: {state['query']}\n\n"]
        for i, q in enumerate(state["questions"][:len(state['interview_responses'])]):
            record.append(f"Question {i+1}: {q}\nYour Response: {state['interview_responses'][i]}\n\n")
        eval_prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an expert interviewer evaluating mock interview responses. Provide detailed feedback on the answers, including strengths, areas for improvement, and an overall score out of 10."),
            ("human", "Evaluate these mock interview responses:\n" + "\n".join(record))
        ])
        eval_response = self.model.invoke(eval_prompt.invoke({}).to_messages())
        record.append(f"AI Evaluation:\n{eval_response.content}\n")
        return ''.join(record)

class ResumeMaker:
    def __init__(self):
        self.model = llm
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled resume expert with extensive experience in crafting resumes tailored for tech roles...'''),
            MessagesPlaceholder('chat_history'),
            ("human", "{input}"),
            MessagesPlaceholder("agent_scratchpad")
        ])
        self.tools = [DuckDuckGoSearchResults()]
        self.agent_executor = AgentExecutor(agent=create_tool_calling_agent(self.model, self.tools, self.prompt), tools=self.tools, verbose=True)
        self.rag_agent = None

    def generate_questions(self, state):
        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a resume expert tasked with gathering details from a user to create a resume. Based on the user's initial query, generate 5 concise, relevant questions to collect essential information (e.g., role, experience, skills, projects, education) for a professional resume. Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        question_response = self.model.invoke(question_prompt.invoke({"query": state['query']}).to_messages())
        return [q.strip() for q in question_response.content.split('\n') if q.strip() and q[0].isdigit()][:5]

    def create_resume(self, state):
        resume_content = f"# Resume\n\nBased on: {state['query']}\n\n"
        for i, answer in state['resume_details'].items():
            resume_content += f"**Detail {i.split('_')[1]}**: {answer}\n\n"
        refine_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a skilled resume expert. Refine this raw resume content into a polished, professional template in Markdown format, incorporating trending keywords and technologies relevant to the specified role.'''),
            ("human", f"Raw resume content:\n{resume_content}")
        ])
        response = self.model.invoke(refine_prompt.invoke({"input": state['query']}).to_messages())
        return response.content

class JobSearch:
    def __init__(self):
        self.model = llm
        self.prompt = ChatPromptTemplate.from_template('''Refactor job search results into .md format.\nContent: {result}''')
        self.tools = DuckDuckGoSearchResults()

    def generate_questions(self, state):
        question_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are a job search assistant. Based on the user's initial query, generate 2 concise clarification question to refine the job search (e.g., location, experience level, job type).'''),
            ("human", f"Initial query: {state['query']}")
        ])
        question_response = self.model.invoke(question_prompt.invoke({"query": state['query']}).to_messages())
        return [question_response.content.strip()]

    def find_jobs(self, state):
        refined_query = f"{state['query']} {state['job_search_details'].get('answer_1', '')}"
        try:
            results = self.tools.run(refined_query)
            if not results:
                results = "No results found for this query."
        except Exception as e:
            results = f"Search failed: {str(e)}."
        chain = self.prompt | self.model
        jobs = chain.invoke({"result": results}).content
        return jobs.replace("```markdown", "").strip()

class LearningResourceAgent:
    def __init__(self):
        self.model = llm
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert in providing learning resources...'''),
            MessagesPlaceholder('chat_history'),
            ("human", "{input}"),
            MessagesPlaceholder("agent_scratchpad")
        ])
        self.tools = [DuckDuckGoSearchResults()]

    def generate_questions(self, state):
        clarification_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert in guiding users to learning resources. Based on the user's initial query, generate 3 concise, relevant clarification questions to tailor the recommendations (e.g., about current knowledge level, preferred format, or goals). Return the questions as a numbered list.'''),
            ("human", f"Initial query: {state['query']}")
        ])
        clarification_response = self.model.invoke(clarification_prompt.invoke({"query": state['query']}).to_messages())
        return [q.strip() for q in clarification_response.content.split('\n') if q.strip() and q[0].isdigit()][:3]

    def tutorial_agent(self, state):
        agent = create_tool_calling_agent(self.model, self.tools, self.prompt)
        agent_executor = AgentExecutor(agent=agent, tools=self.tools, verbose=True)
        response = agent_executor.invoke({"input": state['query'], "chat_history": state.get('chat_history', [])})
        return str(response.get('output', '')).replace("```markdown", "").strip()

    def query_bot(self, state):
        tailored_prompt = ChatPromptTemplate.from_messages([
            ("system", '''You are an expert in providing learning resources. Based on the user's initial query and their clarification answers, generate tailored recommendations for learning sources in Markdown format.'''),
            ("human", f"Initial query: {state['query']}\nClarification 1: {state['learning_details'].get('answer_1', 'N/A')}\nClarification 2: {state['learning_details'].get('answer_2', 'N/A')}\nClarification 3: {state['learning_details'].get('answer_3', 'N/A')}")
        ])
        tailored_response = self.model.invoke(tailored_prompt.invoke({"query": state['query']}).to_messages())
        return tailored_response.content

def save_file(data, filename, format_choice="md"):
    folder_name = "Agent_output"
    os.makedirs(folder_name, exist_ok=True)
    safe_filename = re.sub(r'[<>:"/\\|?*]', '', filename)[:50]
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_path = os.path.join(folder_name, f"{safe_filename}_{timestamp}.{format_choice}")
    if format_choice == "md":
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(data)
    elif format_choice == "docx":
        doc = DocxDocument()
        for line in data.split('\n'):
            if line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('## '): doc.add_heading(line[3:], level=2)
            elif line.strip(): doc.add_paragraph(line)
        doc.save(file_path)
    elif format_choice == "pdf":
        pdf = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for line in data.split('\n'):
            if line.startswith('# '): story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '): story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.strip(): story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 12))
        pdf.build(story)
    return file_path

def extract_resume(file_path):
    if not os.path.exists(file_path):
        return {"status": "error", "message": "File not found"}
    file_extension = os.path.splitext(file_path.lower())[1]
    try:
        if file_extension == ".pdf":
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = "".join(page.extract_text() + "\n" for page in reader.pages)
                return {"status": "success", "content": text}
        elif file_extension in ['.doc', ".docx"]:
            doc = DocxDocument(file_path)
            text = "".join(para.text + "\n" for para in doc.paragraphs)
            return {"status": "success", "content": text}
        return {"status": "error", "message": "Unsupported file format."}
    except Exception as e:
        return {"status": "error", "message": f"Error reading file: {str(e)}"}

def categories(state: State) -> State:
    if state["query"].lower() == "exit":
        return {"category": "exit"}
    prompt = ChatPromptTemplate.from_template(
        "Categorize into one of these: 1: Learn, 2: Resume, 3: Prepare Interview, 4: Job Search, 5: Mock Interview. Give the number only.\nQuery: {query}"
    )
    category = llm.invoke(prompt.invoke({"query": state['query']}).to_messages()).content
    state['chat_history'].append(HumanMessage(content=state['query']))
    return {"category": category}

def handle_learning_resources(state: State) -> State:
    prompt = ChatPromptTemplate.from_template(
        "Categorize as Tutorial or Question.\nQuery: {query}"
    )
    category = llm.invoke(prompt.invoke({"query": state['query']}).to_messages()).content
    return {'category': category}

def job_search(state: State) -> State:
    job_search_instance = JobSearch()
    state["questions"] = job_search_instance.generate_questions(state)
    if state["current_step"] >= len(state["questions"]):
        state["content"] = job_search_instance.find_jobs(state)
        state["response"] = "Job search results generated."
        state['chat_history'].append(AIMessage(content=state["content"]))
    return state

def handle_resume(state: State) -> State:
    resume_maker = ResumeMaker()
    if "grade" in state['query'].lower() and state.get('rag_agent'):
        state["content"] = state['rag_agent'].grade_resume(state)
        state["response"] = "Resume graded."
    else:
        state["questions"] = resume_maker.generate_questions(state)
        if state["current_step"] >= len(state["questions"]):
            state["content"] = resume_maker.create_resume(state)
            state["response"] = "Resume generated."
            state['rag_agent'] = RagAgent()
    state['chat_history'].append(AIMessage(content=state["content"] or "Questions pending"))
    return state

def prepare_interview(state: State) -> State:
    interview_agent = InterviewAgent(state.get('rag_agent'))
    state["questions"] = interview_agent.generate_clarification_questions(state)
    if state["current_step"] >= len(state["questions"]):
        state["content"] = interview_agent.prepare_interview_questions(state)
        state["response"] = "Interview questions generated."
        state['chat_history'].append(AIMessage(content=state["content"]))
    return state

def mock_interview(state: State) -> State:
    interview_agent = InterviewAgent(state.get('rag_agent'))
    if not state["questions"]:
        state["instructions"] = '''**Mock Interview Instructions**:
- You’ll answer 2 clarification questions first to tailor the interview.
- Then, you’ll face 20 questions (5 resume-based if uploaded, 15 technical).
- Answer each question thoughtfully; your responses will be evaluated.
- A timer will start once the actual questions begin.
- You have 60 mins .
- Take your time, but aim to keep answers concise (1-2 minutes each).'''
        state["questions"] = interview_agent.generate_clarification_questions(state)
    elif state["current_step"] >= len(state["questions"]) and len(state["interview_responses"]) < 20:
        state["questions"] = interview_agent.mock_interview_questions(state)
        state["start_time"] = time.time()
        state["current_step"] = 0
        state["instructions"] = None
    if state["current_step"] >= len(state["questions"]):
        state["content"] = interview_agent.evaluate_mock_interview(state)
        state["response"] = "Mock interview completed."
        state['chat_history'].append(AIMessage(content=state["content"]))
    return state

def query_bot(state: State) -> State:
    learning_agent = LearningResourceAgent()
    state["questions"] = learning_agent.generate_questions(state)
    if state["current_step"] >= len(state["questions"]):
        state["content"] = learning_agent.query_bot(state)
        state["response"] = "Learning resources generated."
        state['chat_history'].append(AIMessage(content=state["content"]))
    return state

def tutorial_agent(state: State) -> State:
    learning_agent = LearningResourceAgent()
    state["questions"] = learning_agent.generate_questions(state)
    if state["current_step"] >= len(state["questions"]):
        state["content"] = learning_agent.tutorial_agent(state)
        state["response"] = "Tutorial generated."
        state['chat_history'].append(AIMessage(content=state["content"]))
    return state

def route_query(state: State) -> str:
    if state["category"] == "exit":
        return END
    return {
        "1": "handle_learning_resource",
        "2": "handle_resume_making",
        "3": "prepare_interview",
        "4": "job_search",
        "5": "mock_interview"
    }.get(state["category"], END)

def route_learning(state: State) -> str:
    return "tutorial_agent" if 'Tutorial' in state["category"] else "query_bot"

graph = StateGraph(State)
graph.add_node("categories", categories)
graph.add_node("handle_learning_resource", handle_learning_resources)
graph.add_node("handle_resume_making", handle_resume)
graph.add_node("prepare_interview", prepare_interview)
graph.add_node("job_search", job_search)
graph.add_node("mock_interview", mock_interview)
graph.add_node("tutorial_agent", tutorial_agent)
graph.add_node("query_bot", query_bot)
graph.add_edge(START, "categories")
graph.add_conditional_edges("categories", route_query)
graph.add_conditional_edges("handle_learning_resource", route_learning)
graph.add_edge("handle_resume_making", END)
graph.add_edge("prepare_interview", END)
graph.add_edge("job_search", END)
graph.add_edge("mock_interview", END)
graph.add_edge("tutorial_agent", END)
graph.add_edge("query_bot", END)
app_graph = graph.compile()

def initialize_state() -> State:
    return {
        "query": "",
        "response": None,
        "category": None,
        "chat_history": [],
        "resume_details": {},
        "interview_responses": [],
        "start_time": None,
        "interview_details": {},
        "job_search_details": {},
        "learning_details": {},
        "rag_agent": None,
        "resume_path": None,
        "current_step": 0,
        "questions": [],
        "content": None,
        "instructions": None
    }

state_store: Dict[str, State] = {}

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    # Clear all session data when returning to home
    state_store.clear()
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/start_query", response_class=HTMLResponse)
async def start_query(request: Request, query: str = Form(...), session_id: str = Form(default=None)):
    if not session_id or session_id == "default":
        session_id = str(uuid4())  # Generate a unique session ID
    state_store[session_id] = initialize_state()
    state = state_store[session_id]
    state["query"] = query
    state["chat_history"] = [HumanMessage(content=query)]
    
    results = await asyncio.to_thread(app_graph.invoke, state)
    state.update(results)
    
    if not state["questions"]:
        return templates.TemplateResponse("result.html", {"request": request, "state": state, "session_id": session_id})
    return templates.TemplateResponse("questions.html", {"request": request, "state": state, "session_id": session_id})

@app.post("/submit_answer", response_class=HTMLResponse)
async def submit_answer(request: Request, answer: str = Form(...), session_id: str = Form(...)):
    state = state_store.get(session_id)
    if not state or not state["questions"]:
        return HTMLResponse("Session expired or invalid.")

    step = state["current_step"]
    if state["category"] == "2":
        state["resume_details"][f"answer_{step + 1}"] = answer
    elif state["category"] == "3":
        state["interview_details"][f"answer_{step + 1}"] = answer
    elif state["category"] == "4":
        state["job_search_details"][f"answer_{step + 1}"] = answer
    elif state["category"] == "5":
        state["interview_responses"].append(answer)
    elif state["category"] == "1":
        state["learning_details"][f"answer_{step + 1}"] = answer

    state["current_step"] += 1
    results = await asyncio.to_thread(app_graph.invoke, state)
    state.update(results)
    
    if state["content"]:
        return templates.TemplateResponse("result.html", {"request": request, "state": state, "session_id": session_id})
    return templates.TemplateResponse("questions.html", {"request": request, "state": state, "session_id": session_id})

@app.post("/upload_resume", response_class=HTMLResponse)
async def upload_resume(request: Request, file: UploadFile = File(...), target_role: str = Form(...)):
    session_id = str(uuid4())  # Unique session for resume grading
    state_store[session_id] = initialize_state()
    state = state_store[session_id]
    file_content = await file.read()
    file_extension = os.path.splitext(file.filename.lower())[1]
    temp_path = f"temp_{datetime.now().strftime('%Y%m%d%H%M%S')}{file_extension}"
    
    with open(temp_path, "wb") as f:
        f.write(file_content)
    
    resume_result = extract_resume(temp_path)
    os.remove(temp_path)
    
    if resume_result["status"] == "error":
        state["query"] = f"Upload resume for {target_role}"
        state["response"] =resume_result["message"]
        state["chat_history"] = [HumanMessage(content=state["query"]), AIMessage(content=state["response"])]
        return templates.TemplateResponse("result.html", {"request": request, "state": state, "session_id": session_id})
    
    state["query"] = f"Upload resume for {target_role}"
    state["rag_agent"] = RagAgent(resume_content=resume_result["content"])
    state["resume_path"] = file.filename
    state["content"] = state["rag_agent"].grade_resume(state)
    state["response"] = "Resume graded successfully."
    state["chat_history"] = [HumanMessage(content=state["query"]), AIMessage(content=state["content"])]
    
    return templates.TemplateResponse("result.html", {"request": request, "state": state, "session_id": session_id})

@app.post("/start_mock_interview", response_class=HTMLResponse)
async def start_mock_interview(request: Request, file: UploadFile = File(...), target_role: str = Form(...)):
    session_id = str(uuid4())  # Unique session for mock interview
    state_store[session_id] = initialize_state()
    state = state_store[session_id]
    file_content = await file.read()
    file_extension = os.path.splitext(file.filename.lower())[1]
    temp_path = f"temp_{datetime.now().strftime('%Y%m%d%H%M%S')}{file_extension}"
    
    with open(temp_path, "wb") as f:
        f.write(file_content)
    
    resume_result = extract_resume(temp_path)
    os.remove(temp_path)
    
    if resume_result["status"] == "error":
        state["query"] = f"Conduct a mock interview for {target_role}"
        state["response"] = resume_result["message"]
        state["chat_history"] = [HumanMessage(content=state["query"]), AIMessage(content=state["response"])]
        return templates.TemplateResponse("result.html", {"request": request, "state": state, "session_id": session_id})
    
    state["query"] = f"Conduct a mock interview for {target_role}"
    state["rag_agent"] = RagAgent(resume_content=resume_result["content"])
    state["resume_path"] = file.filename
    state["category"] = "5"
    state["chat_history"] = [HumanMessage(content=state["query"])]
    
    results = await asyncio.to_thread(app_graph.invoke, state)
    state.update(results)
    
    return templates.TemplateResponse("questions.html", {"request": request, "state": state, "session_id": session_id})

@app.post("/save_file")
async def save_file_endpoint(filename: str = Form(...), format_choice: str = Form(...), content: str = Form(...)):
    file_path = save_file(content, filename, format_choice)
    return FileResponse(file_path, filename=os.path.basename(file_path))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
